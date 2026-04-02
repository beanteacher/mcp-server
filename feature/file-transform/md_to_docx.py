"""
Markdown -> DOCX 범용 변환기
맑은 고딕 기반 스타일을 적용하여 임의의 MD 파일을 깔끔한 DOCX로 변환합니다.

지원 문법:
  - 헤딩 (# ~ ####)
  - 불릿/번호 목록 (다단계)
  - 테이블 (헤더 색상, 줄무늬, 리치텍스트)
  - 코드 블록 / 인라인 코드
  - 인용문 (블록쿼트)
  - 인라인 **bold**, `code`
  - 이모지 (Segoe UI Emoji 폰트 자동 적용)

사용법:
    python md_to_docx.py 입력파일.md [출력파일.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_NAME = '맑은 고딕'
FONT_EAST_ASIA = '맑은 고딕'
FONT_EMOJI = 'Segoe UI Emoji'
CODE_FONT = 'Consolas'

_bookmark_id_counter = 0


def _slugify(text):
    """헤딩 텍스트 → 마크다운 스타일 앵커 ID 변환"""
    s = text.lower().strip()
    s = re.sub(r'[*`\[\]()]+', '', s)       # 마크다운 문법 제거
    s = re.sub(r'[^\w\s가-힣-]', '', s)      # 영문/숫자/한글/공백/하이픈만
    s = re.sub(r'\s+', '-', s)               # 공백 → 하이픈
    s = re.sub(r'-+', '-', s).strip('-')     # 중복 하이픈 정리
    return s


def add_bookmark(paragraph, name):
    """문단에 북마크를 추가"""
    global _bookmark_id_counter
    _bookmark_id_counter += 1
    bid = str(_bookmark_id_counter)
    start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bid}" w:name="{name}"/>')
    end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bid}"/>')
    paragraph._element.append(start)
    paragraph._element.append(end)


def add_hyperlink_internal(paragraph, anchor, text, size=10, color=RGBColor(0, 90, 181), bold=False):
    """내부 북마크를 가리키는 하이퍼링크 run 추가"""
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{anchor}"/>')
    run_elem = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr>'
        f'    <w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_EAST_ASIA}"/>'
        f'    <w:sz w:val="{size * 2}"/>'
        f'    <w:color w:val="{color}"/>'
        f'    <w:u w:val="single"/>'
        f'  </w:rPr>'
        f'  <w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


EMOJI_PATTERN = re.compile(
    r'[\U0001F600-\U0001F64F'
    r'\U0001F300-\U0001F5FF'
    r'\U0001F680-\U0001F6FF'
    r'\U0001F900-\U0001F9FF'
    r'\U0001FA00-\U0001FA6F'
    r'\U0001FA70-\U0001FAFF'
    r'\U00002702-\U000027B0'
    r'\U0000FE00-\U0000FE0F'
    r'\u2600-\u26FF'
    r'\u2700-\u27BF'
    r'\u2300-\u23FF'
    r'\u200d'
    r'\u2B50\u2B55'
    r'\u26A0\uFE0F?'
    r']+'
)


def create_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(50, 50, 50)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_EAST_ASIA)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    return doc


def set_font(run, name=FONT_NAME, size=10, color=RGBColor(50, 50, 50), bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_EAST_ASIA)


def set_font_emoji(run, size=10):
    run.font.name = FONT_EMOJI
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_EMOJI)


def add_heading1(doc, text, bookmark_id=None):
    p = doc.add_paragraph()
    if bookmark_id:
        add_bookmark(p, bookmark_id)
    run = p.add_run(text)
    set_font(run, size=18, color=RGBColor(30, 30, 30), bold=True)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_heading2(doc, text, bookmark_id=None):
    p = doc.add_paragraph()
    if bookmark_id:
        add_bookmark(p, bookmark_id)
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(text)
    set_font(run, size=14, color=RGBColor(40, 40, 40), bold=True)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)
    return p


def add_heading3(doc, text, bookmark_id=None):
    p = doc.add_paragraph()
    if bookmark_id:
        add_bookmark(p, bookmark_id)
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text)
    set_font(run, size=12, color=RGBColor(50, 50, 50), bold=True)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_heading4(doc, text, bookmark_id=None):
    p = doc.add_paragraph()
    if bookmark_id:
        add_bookmark(p, bookmark_id)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    set_font(run, size=11, color=RGBColor(60, 60, 60), bold=True)
    p.paragraph_format.space_after = Pt(4)
    return p


def _has_emoji(text):
    return bool(EMOJI_PATTERN.search(text))


def _split_emoji(text):
    segments = []
    last_end = 0
    for m in EMOJI_PATTERN.finditer(text):
        if m.start() > last_end:
            segments.append(('text', text[last_end:m.start()]))
        segments.append(('emoji', m.group()))
        last_end = m.end()
    if last_end < len(text):
        segments.append(('text', text[last_end:]))
    return segments


def add_rich_text(doc, p, text, base_size=10, base_color=RGBColor(50, 50, 50)):
    # Handle markdown links: [text](#anchor) -> internal hyperlink, [text](url) -> text only
    link_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    parts = []
    last_end = 0
    for m in link_pattern.finditer(text):
        if m.start() > last_end:
            parts.append(('text', text[last_end:m.start()]))
        link_text, href = m.group(1), m.group(2)
        if href.startswith('#'):
            parts.append(('link', link_text, href[1:]))  # internal anchor
        else:
            parts.append(('text', link_text))  # external link → plain text
        last_end = m.end()
    if last_end < len(text):
        parts.append(('text', text[last_end:]))

    for part in parts:
        if part[0] == 'link':
            add_hyperlink_internal(p, part[2], part[1], size=base_size)
        else:
            segment = part[1]
            _add_rich_segment(p, segment, base_size, base_color)
    return


def _add_rich_segment(p, text, base_size=10, base_color=RGBColor(50, 50, 50)):
    """bold/code 인라인 처리 (add_rich_text에서 분리)"""
    pattern = re.compile(r'(\*\*(.+?)\*\*|`(.+?)`)')
    last_end = 0

    for m in pattern.finditer(text):
        if m.start() > last_end:
            _add_text_with_emoji(p, text[last_end:m.start()], size=base_size, color=base_color)
        if m.group(2):
            _add_text_with_emoji(p, m.group(2), size=base_size, color=base_color, bold=True)
        elif m.group(3):
            run = p.add_run(m.group(3))
            set_font(run, name=CODE_FONT, size=max(base_size - 1, 8), color=RGBColor(180, 50, 50))
        last_end = m.end()

    if last_end < len(text):
        _add_text_with_emoji(p, text[last_end:], size=base_size, color=base_color)


def _add_text_with_emoji(p, text, size=10, color=RGBColor(50, 50, 50), bold=False):
    if not _has_emoji(text):
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=bold)
        return

    for seg_type, seg_text in _split_emoji(text):
        run = p.add_run(seg_text)
        if seg_type == 'emoji':
            set_font_emoji(run, size=size)
        else:
            set_font(run, size=size, color=color, bold=bold)


def add_body(doc, text):
    p = doc.add_paragraph()
    add_rich_text(doc, p, text)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="8" w:color="B0B0B0"/></w:pBdr>')
    pPr.append(pBdr)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shd)
    add_rich_text(doc, p, text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, name=CODE_FONT, size=9, color=RGBColor(60, 60, 60))
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    base_indent = 0.8
    indent = base_indent + level * 0.8
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(2)

    if level == 0:
        marker = "\u2022 "
        color = RGBColor(50, 50, 50)
    elif level == 1:
        marker = "- "
        color = RGBColor(100, 100, 100)
    else:
        marker = "\u25E6 "
        color = RGBColor(120, 120, 120)

    run = p.add_run(marker)
    set_font(run, color=color)
    add_rich_text(doc, p, text)
    return p


def add_numbered(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{num}. ")
    set_font(run)
    add_rich_text(doc, p, text)
    return p


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_font_header(cell, text, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    _add_text_with_emoji(p, text, size=size, color=RGBColor(255, 255, 255), bold=True)


def set_cell_font_body(cell, text, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    add_rich_text(None, p, text, base_size=size)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "373F51")
        set_cell_font_header(cell, h, size=9)

    for r_idx, row_data in enumerate(rows):
        bg = "F8F9FA" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_shading(cell, bg)
            set_cell_font_body(cell, cell_text, size=9)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    doc.add_paragraph()


# ---- Markdown parser ----

def parse_table(lines, start_idx):
    header_line = lines[start_idx].strip()
    headers = [c.strip() for c in header_line.strip('|').split('|')]

    sep_idx = start_idx + 1
    if sep_idx < len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[sep_idx]):
        sep_idx += 1

    rows = []
    idx = sep_idx
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith('|'):
            break
        row = [c.strip() for c in line.strip('|').split('|')]
        while len(row) < len(headers):
            row.append("")
        rows.append(row[:len(headers)])
        idx += 1

    return headers, rows, idx


def convert_md_to_docx(md_path, docx_path):
    md_text = Path(md_path).read_text(encoding='utf-8')
    lines = md_text.split('\n')
    doc = create_doc()

    idx = 0
    in_code_block = False
    code_lines = []

    while idx < len(lines):
        line = lines[idx]

        # code block
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
                idx += 1
                continue
            else:
                in_code_block = True
                idx += 1
                continue

        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue

        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        # horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            idx += 1
            continue

        # heading: \s* makes the space after # optional so "#1. Title" works
        heading_match = re.match(r'^(#{1,4})\s*(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            bid = _slugify(text)
            if level == 1:
                add_heading1(doc, text, bookmark_id=bid)
            elif level == 2:
                add_heading2(doc, text, bookmark_id=bid)
            elif level == 3:
                add_heading3(doc, text, bookmark_id=bid)
            elif level == 4:
                add_heading4(doc, text, bookmark_id=bid)
            idx += 1
            continue

        # table
        if stripped.startswith('|') and idx + 1 < len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[idx + 1]):
            headers, rows, idx = parse_table(lines, idx)
            clean_headers = [re.sub(r'\*\*(.+?)\*\*', r'\1', h) for h in headers]
            add_table(doc, clean_headers, rows)
            continue

        # blockquote
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            while idx + 1 < len(lines) and lines[idx + 1].strip().startswith('>'):
                idx += 1
                next_line = lines[idx].strip().lstrip('>').strip()
                if next_line:
                    quote_text += '\n' + next_line
            add_quote(doc, quote_text)
            idx += 1
            continue

        # numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            add_numbered(doc, num_match.group(1), num_match.group(2))
            idx += 1
            continue

        # bullet list (indent-aware)
        bullet_match = re.match(r'^(\s*)([-*])\s+(.+)', line)
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            text = bullet_match.group(3)
            if indent_spaces < 2:
                level = 0
            elif indent_spaces < 4:
                level = 1
            else:
                level = min(indent_spaces // 2, 3)
            add_bullet(doc, text, level=level)
            idx += 1
            continue

        # plain text
        add_body(doc, stripped)
        idx += 1

    doc.save(docx_path)
    print(f"OK:{docx_path}")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py input.md [output.docx]")
        sys.exit(1)

    md_file = sys.argv[1]
    if len(sys.argv) >= 3:
        docx_file = sys.argv[2]
    else:
        docx_file = str(Path(md_file).with_suffix('.docx'))

    convert_md_to_docx(md_file, docx_file)
