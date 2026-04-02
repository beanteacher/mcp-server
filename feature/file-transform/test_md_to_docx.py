"""md_to_docx.py 단위 테스트"""
import os
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from md_to_docx import (
    _slugify,
    _has_emoji,
    _split_emoji,
    add_bookmark,
    add_heading1,
    add_heading2,
    add_heading3,
    add_heading4,
    add_hyperlink_internal,
    add_rich_text,
    add_body,
    add_bullet,
    add_numbered,
    add_quote,
    add_code_block,
    add_table,
    create_doc,
    parse_table,
    convert_md_to_docx,
)


# ── _slugify ──


class TestSlugify:
    def test_basic_korean(self):
        assert _slugify("1. 서론") == "1-서론"

    def test_english(self):
        assert _slugify("Hello World") == "hello-world"

    def test_mixed_korean_english(self):
        assert _slugify("KVKK 법제 및 국외 데이터 반출 규제 심층 분석") == "kvkk-법제-및-국외-데이터-반출-규제-심층-분석"

    def test_strips_markdown_syntax(self):
        assert _slugify("**bold** and `code`") == "bold-and-code"

    def test_strips_parentheses_brackets(self):
        assert _slugify("굴삭기(건설장비) 브랜드별") == "굴삭기건설장비-브랜드별"

    def test_collapses_multiple_hyphens(self):
        assert _slugify("a  -  b") == "a-b"

    def test_strips_leading_trailing_hyphens(self):
        assert _slugify(" - hello - ") == "hello"

    def test_empty_string(self):
        assert _slugify("") == ""

    def test_numbered_heading(self):
        assert _slugify("3. KVKK 법제 및 국외 데이터 반출 규제 심층 분석") == "3-kvkk-법제-및-국외-데이터-반출-규제-심층-분석"


# ── emoji helpers ──


class TestEmojiHelpers:
    def test_has_emoji_true(self):
        assert _has_emoji("hello 🚀 world")

    def test_has_emoji_false(self):
        assert not _has_emoji("hello world")

    def test_split_emoji(self):
        segments = _split_emoji("hello 🚀 world")
        assert segments[0] == ("text", "hello ")
        assert segments[1] == ("emoji", "🚀")
        assert segments[2] == ("text", " world")

    def test_split_no_emoji(self):
        segments = _split_emoji("plain text")
        assert segments == [("text", "plain text")]


# ── bookmark ──


class TestBookmark:
    def test_bookmark_added_to_paragraph(self):
        doc = create_doc()
        p = doc.add_paragraph("test")
        add_bookmark(p, "my-bookmark")

        xml = p._element.xml
        assert "bookmarkStart" in xml
        assert "bookmarkEnd" in xml
        assert 'w:name="my-bookmark"' in xml


# ── headings with bookmark ──


class TestHeadings:
    def test_heading1_text(self):
        doc = create_doc()
        p = add_heading1(doc, "제목1")
        assert p.runs[0].text == "제목1"
        assert p.runs[0].bold is True

    def test_heading2_with_bookmark(self):
        doc = create_doc()
        p = add_heading2(doc, "2. 시장 개요", bookmark_id="2-시장-개요")
        xml = p._element.xml
        assert 'w:name="2-시장-개요"' in xml
        assert p.runs[0].text == "2. 시장 개요"

    def test_heading3_no_bookmark(self):
        doc = create_doc()
        p = add_heading3(doc, "3.1 소제목")
        assert "bookmarkStart" not in p._element.xml

    def test_heading4_with_bookmark(self):
        doc = create_doc()
        p = add_heading4(doc, "세부항목", bookmark_id="세부항목")
        assert 'w:name="세부항목"' in p._element.xml


# ── internal hyperlink ──


class TestHyperlinkInternal:
    def test_hyperlink_anchor(self):
        doc = create_doc()
        p = doc.add_paragraph()
        add_hyperlink_internal(p, "1-서론", "서론")

        xml = p._element.xml
        assert 'w:anchor="1-서론"' in xml
        assert "서론" in xml

    def test_hyperlink_has_underline(self):
        doc = create_doc()
        p = doc.add_paragraph()
        add_hyperlink_internal(p, "anchor", "link text")

        xml = p._element.xml
        assert 'w:val="single"' in xml  # underline


# ── add_rich_text (link handling) ──


class TestAddRichText:
    def test_internal_link_creates_hyperlink(self):
        doc = create_doc()
        p = doc.add_paragraph()
        add_rich_text(doc, p, "목차: [서론](#1-서론)")

        xml = p._element.xml
        assert 'w:anchor="1-서론"' in xml
        assert "서론" in xml

    def test_external_link_becomes_plain_text(self):
        doc = create_doc()
        p = doc.add_paragraph()
        add_rich_text(doc, p, "[Google](https://google.com)")

        text = "".join(r.text for r in p.runs)
        assert "Google" in text
        assert "https" not in p._element.xml

    def test_bold_inline(self):
        doc = create_doc()
        p = doc.add_paragraph()
        add_rich_text(doc, p, "일반 **굵게** 텍스트")

        bold_runs = [r for r in p.runs if r.bold]
        assert any("굵게" in r.text for r in bold_runs)

    def test_code_inline(self):
        doc = create_doc()
        p = doc.add_paragraph()
        add_rich_text(doc, p, "코드 `hello()` 입니다")

        code_runs = [r for r in p.runs if r.font.name == "Consolas"]
        assert any("hello()" in r.text for r in code_runs)

    def test_plain_text_no_links(self):
        doc = create_doc()
        p = doc.add_paragraph()
        add_rich_text(doc, p, "링크 없는 일반 텍스트")

        text = "".join(r.text for r in p.runs)
        assert text == "링크 없는 일반 텍스트"


# ── document elements ──


class TestDocumentElements:
    def test_add_body(self):
        doc = create_doc()
        p = add_body(doc, "본문 텍스트")
        assert len(p.runs) > 0

    def test_add_bullet_levels(self):
        doc = create_doc()
        p0 = add_bullet(doc, "레벨0", level=0)
        p1 = add_bullet(doc, "레벨1", level=1)
        assert p0.paragraph_format.left_indent < p1.paragraph_format.left_indent

    def test_add_numbered(self):
        doc = create_doc()
        p = add_numbered(doc, "1", "첫번째 항목")
        assert p.runs[0].text == "1. "

    def test_add_quote(self):
        doc = create_doc()
        p = add_quote(doc, "인용문입니다")
        assert p.paragraph_format.left_indent is not None

    def test_add_code_block(self):
        doc = create_doc()
        p = add_code_block(doc, "const x = 1;")
        assert p.runs[0].text == "const x = 1;"
        assert p.runs[0].font.name == "Consolas"

    def test_add_table(self):
        doc = create_doc()
        add_table(doc, ["이름", "값"], [["A", "1"], ["B", "2"]])
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 3  # header + 2 rows


# ── parse_table ──


class TestParseTable:
    def test_basic_table(self):
        lines = [
            "| 이름 | 값 |",
            "| --- | --- |",
            "| A | 1 |",
            "| B | 2 |",
            "",
        ]
        headers, rows, end_idx = parse_table(lines, 0)
        assert headers == ["이름", "값"]
        assert len(rows) == 2
        assert rows[0] == ["A", "1"]
        assert end_idx == 4

    def test_table_with_missing_cells(self):
        lines = [
            "| A | B | C |",
            "| --- | --- | --- |",
            "| 1 | 2 |",
        ]
        headers, rows, _ = parse_table(lines, 0)
        assert len(rows[0]) == 3  # padded


# ── convert_md_to_docx (통합 테스트) ──


class TestConvertMdToDocx:
    def _convert(self, md_content: str) -> Document:
        with tempfile.TemporaryDirectory() as tmpdir:
            md_path = os.path.join(tmpdir, "test.md")
            docx_path = os.path.join(tmpdir, "test.docx")
            Path(md_path).write_text(md_content, encoding="utf-8")
            convert_md_to_docx(md_path, docx_path)
            return Document(docx_path)

    def test_heading_bookmark_and_toc_link(self):
        md = (
            "## 목차\n\n"
            "1. [서론](#1-서론)\n\n"
            "---\n\n"
            "## 1. 서론\n\n"
            "본문입니다.\n"
        )
        doc = self._convert(md)
        full_xml = doc.element.xml

        # 헤딩에 북마크가 있어야 함
        assert 'w:name="1-서론"' in full_xml
        # 목차 링크가 하이퍼링크로 변환되어야 함
        assert 'w:anchor="1-서론"' in full_xml

    def test_multiple_headings_get_bookmarks(self):
        md = "# Title\n\n## Section A\n\n## Section B\n"
        doc = self._convert(md)
        xml = doc.element.xml
        assert 'w:name="title"' in xml
        assert 'w:name="section-a"' in xml
        assert 'w:name="section-b"' in xml

    def test_code_block(self):
        md = "```\nconst x = 1;\n```\n"
        doc = self._convert(md)
        texts = [p.text for p in doc.paragraphs]
        assert any("const x = 1;" in t for t in texts)

    def test_blockquote(self):
        md = "> 인용문 테스트\n"
        doc = self._convert(md)
        texts = [p.text for p in doc.paragraphs]
        assert any("인용문 테스트" in t for t in texts)

    def test_table_conversion(self):
        md = "| 헤더1 | 헤더2 |\n| --- | --- |\n| A | B |\n"
        doc = self._convert(md)
        assert len(doc.tables) == 1

    def test_bullet_list(self):
        md = "- 항목1\n- 항목2\n  - 하위항목\n"
        doc = self._convert(md)
        texts = [p.text for p in doc.paragraphs]
        assert any("항목1" in t for t in texts)

    def test_horizontal_rule_ignored(self):
        md = "텍스트\n\n---\n\n다음 텍스트\n"
        doc = self._convert(md)
        texts = [p.text for p in doc.paragraphs]
        assert not any(t.strip() == "---" for t in texts)

    def test_empty_md(self):
        doc = self._convert("")
        assert len(doc.paragraphs) >= 0  # no crash
